"""Services supporting submission handling: files, reviewers and letters."""

from __future__ import annotations

import logging
import re
from datetime import timedelta
from typing import Any

from django.conf import settings
from django.core.exceptions import ValidationError
from django.db.models import Avg, Count, Q
from django.utils import timezone
from django.utils.translation import gettext as _

from apps.accounts.models import Role, User
from apps.submissions.models import (
    Review,
    ReviewAssignment,
    ReviewRound,
    Submission,
    SubmissionFile,
)

logger = logging.getLogger(__name__)

#: Allowed upload extensions per file kind (SPEC §11).
ALLOWED_EXTENSIONS: dict[str, tuple[str, ...]] = {
    SubmissionFile.Kind.MANUSCRIPT_ANON: ("pdf", "docx", "doc", "odt", "tex"),
    SubmissionFile.Kind.TITLE_PAGE: ("pdf", "docx", "doc", "odt"),
    SubmissionFile.Kind.FIGURES: ("pdf", "zip", "png", "jpg", "jpeg", "svg", "docx"),
    SubmissionFile.Kind.DATA: ("zip", "csv", "xlsx", "pdf", "txt"),
    SubmissionFile.Kind.SUPPLEMENTARY: ("pdf", "zip", "docx", "csv", "xlsx", "png", "jpg"),
    SubmissionFile.Kind.REVISION: ("pdf", "docx", "doc", "odt", "tex"),
    SubmissionFile.Kind.RESPONSE: ("pdf", "docx", "doc", "odt", "txt"),
    SubmissionFile.Kind.COPYEDITED: ("pdf", "docx", "odt"),
    SubmissionFile.Kind.PROOF: ("pdf",),
    SubmissionFile.Kind.FINAL: ("pdf", "docx", "xml"),
}

#: Extension → expected MIME prefix, checked with python-magic when available.
MIME_HINTS: dict[str, tuple[str, ...]] = {
    "pdf": ("application/pdf",),
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
    ),
    "doc": ("application/msword", "application/x-ole-storage"),
    "odt": ("application/vnd.oasis.opendocument.text", "application/zip"),
    "zip": ("application/zip",),
    "csv": ("text/plain", "text/csv", "application/csv"),
    "xlsx": (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/zip",
    ),
    "png": ("image/png",),
    "jpg": ("image/jpeg",),
    "jpeg": ("image/jpeg",),
    "svg": ("image/svg+xml", "text/plain", "text/xml"),
    "xml": ("text/xml", "application/xml", "text/plain"),
    "tex": ("text/plain", "text/x-tex"),
    "txt": ("text/plain",),
}


def validate_upload(uploaded_file, kind: str) -> None:
    """Validate size, extension and sniffed MIME type of an upload."""
    max_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if uploaded_file.size > max_bytes:
        raise ValidationError(
            _("The file is larger than %(limit)s MB.") % {"limit": settings.MAX_UPLOAD_SIZE_MB}
        )

    name = uploaded_file.name or ""
    extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    allowed = ALLOWED_EXTENSIONS.get(kind, ("pdf", "docx"))
    if extension not in allowed:
        raise ValidationError(
            _("Allowed file types for this slot: %(types)s.") % {"types": ", ".join(allowed)}
        )

    head = uploaded_file.read(4096)
    uploaded_file.seek(0)
    detected = _sniff_mime(head)
    if detected:
        expected = MIME_HINTS.get(extension, ())
        if expected and not any(detected.startswith(e) for e in expected):
            raise ValidationError(
                _("The file content (%(detected)s) does not match its extension.")
                % {"detected": detected}
            )
    scan_for_viruses(head)


def _sniff_mime(head: bytes) -> str:
    """Detect the MIME type of a byte prefix, returning '' when unavailable."""
    try:  # pragma: no cover - libmagic may be absent on some hosts
        import magic

        return magic.from_buffer(head, mime=True) or ""
    except Exception:
        return ""


def scan_for_viruses(payload: bytes) -> None:
    """Optionally scan an upload with ClamAV when ``CLAMAV_HOST`` is set."""
    host = settings.CLAMAV_HOST
    if not host:
        return
    try:  # pragma: no cover - optional infrastructure
        import socket

        with socket.create_connection((host, 3310), timeout=5) as sock:
            sock.sendall(b"zINSTREAM\0")
            sock.sendall(len(payload).to_bytes(4, "big") + payload + b"\0\0\0\0")
            response = sock.recv(256).decode(errors="ignore")
        if "FOUND" in response:
            raise ValidationError(_("The uploaded file failed the virus scan."))
    except ValidationError:
        raise
    except Exception:
        logger.warning("ClamAV scan skipped: host %s unreachable", host)


def count_words(uploaded_file) -> int:
    """Estimate the manuscript word count from a DOCX or PDF upload."""
    name = (uploaded_file.name or "").lower()
    try:
        uploaded_file.seek(0)
        if name.endswith(".docx"):
            from docx import Document

            document = Document(uploaded_file)
            words = sum(len(p.text.split()) for p in document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        words += len(cell.text.split())
            return words
        if name.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(uploaded_file)
            return sum(len((page.extract_text() or "").split()) for page in reader.pages)
    except Exception:  # pragma: no cover - malformed documents
        logger.info("Could not count words in %s", name)
    finally:
        uploaded_file.seek(0)
    return 0


def page_count(uploaded_file) -> int:
    """Number of pages in a PDF upload (0 for other formats)."""
    if not (uploaded_file.name or "").lower().endswith(".pdf"):
        return 0
    try:
        uploaded_file.seek(0)
        from pypdf import PdfReader

        return len(PdfReader(uploaded_file).pages)
    except Exception:  # pragma: no cover
        return 0
    finally:
        uploaded_file.seek(0)


#: Document-information keys that can identify an author or their institution.
IDENTIFYING_PDF_KEYS: tuple[str, ...] = (
    "/Author",
    "/Creator",
    "/Producer",
    "/Title",
    "/Subject",
    "/Keywords",
    "/Company",
    "/SourceModified",
)


def strip_pdf_metadata(file_field) -> bool:
    """Remove identifying metadata from a PDF in place (reviewer anonymity).

    Both the document information dictionary and the XMP packet are cleared.
    ``pikepdf`` writes its own ``/Producer`` when saving, which carries no
    information about the author, so that key may reappear.
    """
    try:
        import pikepdf

        path = file_field.path
        with pikepdf.open(path, allow_overwriting_input=True) as pdf:
            # ``docinfo`` is a PDF dictionary: delete each key rather than
            # calling ``clear()``, which pikepdf only defines for arrays.
            for key in list(pdf.docinfo.keys()):
                del pdf.docinfo[key]
            with pdf.open_metadata() as meta:
                for key in list(meta.keys()):
                    del meta[key]
            pdf.save(path)
        return True
    except Exception:  # pragma: no cover - non-PDF or storage without paths
        logger.exception("Could not strip PDF metadata from %s", getattr(file_field, "name", "?"))
        return False


# ---------------------------------------------------------------------------
# Reviewer management
# ---------------------------------------------------------------------------
def find_reviewers(
    *,
    query: str = "",
    jel_codes: list[str] | None = None,
    exclude_users: list[int] | None = None,
    limit: int = 25,
) -> list[dict[str, Any]]:
    """Rank candidate reviewers by expertise, workload and past performance."""
    queryset = (
        User.objects.filter(is_reviewer=True, is_active=True)
        .select_related("profile")
        .exclude(pk__in=exclude_users or [])
    )
    if query:
        queryset = queryset.filter(
            Q(first_name__icontains=query)
            | Q(last_name__icontains=query)
            | Q(email__icontains=query)
            | Q(profile__expertise__icontains=query)
            | Q(profile__affiliation__icontains=query)
        )
    if jel_codes:
        queryset = queryset.filter(profile__jel_codes__code__in=jel_codes).distinct()

    queryset = queryset.annotate(
        active_reviews=Count(
            "review_assignments",
            filter=Q(
                review_assignments__status__in=[
                    ReviewAssignment.Status.INVITED,
                    ReviewAssignment.Status.ACCEPTED,
                ]
            ),
            distinct=True,
        ),
        completed_reviews=Count(
            "review_assignments",
            filter=Q(review_assignments__status=ReviewAssignment.Status.SUBMITTED),
            distinct=True,
        ),
        quality=Avg("review_assignments__review__quality_rating"),
    ).order_by("active_reviews", "-completed_reviews")[:limit]

    return [
        {
            "user": user,
            "active_reviews": user.active_reviews,
            "completed_reviews": user.completed_reviews,
            "quality": round(user.quality, 1) if user.quality else None,
            "average_days": getattr(user.profile, "average_review_days", None)
            if hasattr(user, "profile")
            else None,
            "expertise": user.profile.expertise_list if hasattr(user, "profile") else [],
        }
        for user in queryset
    ]


def invite_reviewer(
    round_obj: ReviewRound,
    reviewer: User,
    *,
    invited_by: User,
    due_days: int | None = None,
) -> ReviewAssignment:
    """Create a review assignment and queue the invitation e-mail."""
    if ReviewAssignment.objects.filter(round=round_obj, reviewer=reviewer).exists():
        raise ValidationError(_("This reviewer has already been invited to this round."))
    if reviewer.pk == round_obj.submission.submitter_id:
        raise ValidationError(_("The submitting author cannot review their own manuscript."))
    if round_obj.submission.authors.filter(email__iexact=reviewer.email).exists():
        raise ValidationError(_("This person is an author of the manuscript."))

    assignment = ReviewAssignment.objects.create(
        round=round_obj,
        reviewer=reviewer,
        invited_by=invited_by,
        due_at=timezone.now() + timedelta(days=due_days or settings.REVIEW_DUE_DAYS),
    )
    from apps.submissions.tasks import send_reviewer_invitation

    send_reviewer_invitation.delay(assignment.pk)
    return assignment


def invite_reviewer_by_email(
    round_obj: ReviewRound,
    email: str,
    *,
    invited_by: User,
    first_name: str = "",
    last_name: str = "",
) -> ReviewAssignment:
    """Invite someone who has no account yet, creating an inactive user."""
    email = email.strip().lower()
    user = User.objects.filter(email=email).first()
    created = False
    if user is None:
        user = User.objects.create_user(
            email=email,
            password=None,
            first_name=first_name,
            last_name=last_name,
            is_reviewer=True,
        )
        user.set_unusable_password()
        user.save(update_fields=["password"])
        created = True
    if not user.is_reviewer:
        user.is_reviewer = True
        user.save(update_fields=["is_reviewer"])
    _ensure_group(user, Role.REVIEWER)
    assignment = invite_reviewer(round_obj, user, invited_by=invited_by)
    if created:
        from apps.submissions.tasks import send_account_invitation

        send_account_invitation.delay(user.pk)
    return assignment


def _ensure_group(user: User, role: str) -> None:
    """Add the user to a role group, creating the group when needed."""
    from django.contrib.auth.models import Group

    group, _created = Group.objects.get_or_create(name=role)
    user.groups.add(group)


def submit_review(assignment: ReviewAssignment, review: Review) -> Review:
    """Finalise a review and advance the round when everything is in."""
    review.is_draft = False
    review.submitted_at = timezone.now()
    if review.attachment:
        review.is_anonymised = strip_pdf_metadata(review.attachment)
    review.save()

    assignment.status = ReviewAssignment.Status.SUBMITTED
    assignment.completed_at = timezone.now()
    assignment.save(update_fields=["status", "completed_at", "updated_at"])

    profile = getattr(assignment.reviewer, "profile", None)
    if profile is not None:
        profile.reviews_completed += 1
        days = (assignment.completed_at - assignment.invited_at).days
        previous = profile.average_review_days
        profile.average_review_days = days if previous is None else round((previous + days) / 2, 1)
        profile.save(update_fields=["reviews_completed", "average_review_days", "updated_at"])

    _maybe_close_round(assignment.round)
    from apps.submissions.tasks import send_review_thanks

    send_review_thanks.delay(assignment.pk)
    return review


def _maybe_close_round(round_obj: ReviewRound) -> None:
    """Move the submission to ``awaiting_decision`` when reviews are complete."""
    from apps.submissions import workflow
    from apps.submissions.models import SubmissionStatus

    submission = round_obj.submission
    if submission.status != SubmissionStatus.UNDER_REVIEW:
        return
    active = round_obj.assignments.exclude(
        status__in=[ReviewAssignment.Status.DECLINED, ReviewAssignment.Status.CANCELLED]
    )
    submitted = active.filter(status=ReviewAssignment.Status.SUBMITTED).count()
    if submitted >= settings.MIN_REVIEWERS_PER_ROUND and submitted == active.count():
        editor = submission.assigned_editor
        if editor is not None and workflow.can(submission, "reviews_complete", editor):
            workflow.perform(submission, "reviews_complete", editor)
        else:
            submission.status = SubmissionStatus.AWAITING_DECISION
            submission.last_activity_at = timezone.now()
            submission.save(update_fields=["status", "last_activity_at", "updated_at"])
            workflow.system_message(
                submission, str(_("All reviews received; awaiting the editorial decision."))
            )


def build_decision_letter(submission: Submission, decision: str) -> str:
    """Render a starting decision letter merging the reviews of this round."""
    from apps.core.services import get_site_settings

    site = get_site_settings()
    reviews = Review.objects.filter(
        assignment__round=submission.latest_round, is_draft=False
    ).order_by("submitted_at")

    intro = {
        "accept": _("I am pleased to inform you that your manuscript has been accepted."),
        "minor_revision": _("Your manuscript requires minor revisions before it can be accepted."),
        "major_revision": _(
            "Your manuscript requires major revisions before it can be considered further."
        ),
        "reject": _("After careful consideration, we cannot accept your manuscript."),
        "desk_reject": _(
            "After an initial editorial assessment, we cannot proceed with your manuscript."
        ),
        "resubmit": _(
            "We cannot accept the manuscript in its present form, but we would welcome a resubmission."
        ),
    }.get(decision, "")

    lines = [
        _("Dear Author,"),
        "",
        str(intro),
        "",
        _("Manuscript: %(title)s (%(reference)s)")
        % {"title": submission.title, "reference": submission.reference or "-"},
        "",
    ]
    for index, review in enumerate(reviews, start=1):
        lines.append(str(_("Reviewer %(n)d") % {"n": index}))
        lines.append("")
        lines.append(review.comments_to_authors.strip())
        lines.append("")
    lines.extend(
        [
            _("Yours sincerely,"),
            _("Editorial Office"),
            site.journal_name,
        ]
    )
    return "\n".join(str(line) for line in lines)


def normalise_keywords(text: str) -> list[str]:
    """Split a comma-separated keyword string into a clean list."""
    return [re.sub(r"\s+", " ", part).strip() for part in (text or "").split(",") if part.strip()]
